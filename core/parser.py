"""
Resume & JD Parser — extracts raw text from PDF/DOCX files
and parses structured fields (name, email, phone, skills, etc.)
"""

import re
import io
import logging

logger = logging.getLogger(__name__)


# ─── Text Extraction ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path_or_obj):
    """Extract text from a PDF file. Tries pdfplumber first, PyMuPDF as fallback."""
    text = ""

    # Handle Django FieldFile / file-like objects
    if hasattr(file_path_or_obj, 'read'):
        file_bytes = file_path_or_obj.read()
        if hasattr(file_path_or_obj, 'seek'):
            file_path_or_obj.seek(0)
    else:
        with open(file_path_or_obj, 'rb') as f:
            file_bytes = f.read()

    # Try pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # Fallback: PyMuPDF
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        if text.strip():
            return text.strip()
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}")

    logger.error("Could not extract text from PDF")
    return ""


def extract_text_from_docx(file_path_or_obj):
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        from docx import Document
        if hasattr(file_path_or_obj, 'read'):
            file_bytes = file_path_or_obj.read()
            if hasattr(file_path_or_obj, 'seek'):
                file_path_or_obj.seek(0)
            doc = Document(io.BytesIO(file_bytes))
        else:
            doc = Document(file_path_or_obj)

        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        logger.error(f"python-docx failed: {e}")
    return text.strip()


def extract_text(file_field):
    """Auto-detect file type and extract text."""
    name = file_field.name.lower()
    if name.endswith('.pdf'):
        return extract_text_from_pdf(file_field)
    elif name.endswith('.docx') or name.endswith('.doc'):
        return extract_text_from_docx(file_field)
    else:
        # Try reading as plain text
        try:
            return file_field.read().decode('utf-8', errors='ignore')
        except Exception:
            return ""


# ─── Field Parsers ────────────────────────────────────────────────────────────

def extract_email(text):
    emails = re.findall(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
    return emails[0] if emails else ""


def extract_phone(text):
    phones = re.findall(
        r'(?:\+91[-\s]?)?(?:\(?\d{3,5}\)?[-\s]?)?\d{3}[-\s]?\d{4}', text
    )
    return phones[0].strip() if phones else ""


def extract_name(text):
    """
    Heuristic: The candidate name is usually on one of the first 3 lines
    and looks like 'Firstname Lastname' — no numbers, short line.
    """
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        # Skip lines with email, phone, URLs or common headers
        if re.search(r'[@://|•|resume|curriculum|vitae|\d{5,}]', line, re.IGNORECASE):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
            return line
    return lines[0] if lines else ""


def extract_section(text, section_keywords, next_section_keywords=None):
    """
    Extract a section of text by finding lines that match section headers.
    """
    lines = text.split('\n')
    in_section = False
    section_text = []
    next_kws = next_section_keywords or []

    for line in lines:
        line_lower = line.lower().strip()
        if any(kw in line_lower for kw in section_keywords) and len(line.strip()) < 50:
            in_section = True
            continue
        if in_section:
            if next_kws and any(kw in line_lower for kw in next_kws) and len(line.strip()) < 50:
                break
            section_text.append(line)

    return '\n'.join(section_text).strip()


def parse_resume(text):
    """
    Parse a resume text into structured fields.
    Returns a dict: name, email, phone, skills_text, education_text,
    experience_text, projects_text
    """
    if not text:
        return {}

    return {
        'name': extract_name(text),
        'email': extract_email(text),
        'phone': extract_phone(text),
        'skills_text': extract_section(
            text,
            ['skills', 'technical skills', 'core competencies', 'technologies', 'tools'],
            ['experience', 'education', 'projects', 'work', 'employment', 'certifications']
        ),
        'education_text': extract_section(
            text,
            ['education', 'academic', 'qualification', 'degree'],
            ['experience', 'skills', 'projects', 'work', 'employment']
        ),
        'experience_text': extract_section(
            text,
            ['experience', 'work experience', 'employment', 'professional experience', 'work history'],
            ['education', 'skills', 'projects', 'certifications', 'achievements']
        ),
        'projects_text': extract_section(
            text,
            ['projects', 'personal projects', 'academic projects', 'key projects'],
            ['experience', 'education', 'skills', 'certifications', 'achievements']
        ),
    }


def parse_job_description(text):
    """
    Parse a job description text to extract key fields.
    """
    return {
        'skills_text': extract_section(
            text,
            ['requirements', 'required skills', 'skills', 'qualifications', 'must have', 'technologies'],
            ['responsibilities', 'about', 'benefits', 'salary']
        ),
        'experience_text': extract_section(
            text,
            ['experience', 'years of experience'],
            ['skills', 'education', 'responsibilities']
        ),
        'education_text': extract_section(
            text,
            ['education', 'qualification', 'degree'],
            ['skills', 'experience', 'responsibilities']
        ),
    }


def detect_resume_fraud(text, experience_text=None):
    """
    Analyzes resume text for common fraud/gaming heuristics.
    Returns: (fraud_score, list_of_warning_flags)
    """
    if not text:
        return 0, []

    fraud_score = 0
    warning_flags = []
    text_lower = text.lower()

    # 1. Keyword Stuffing Check
    from core.nlp import extract_skills
    skills = extract_skills(text)
    for skill in skills:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        count = len(re.findall(pattern, text_lower))
        if count > 7:
            warning_flags.append(f"Keyword Stuffing: Skill '{skill}' is repeated {count} times (suspiciously high).")
            fraud_score += 15

    # 2. Unrealistic Experience Checks
    from core.nlp import extract_years_of_experience
    total_years = extract_years_of_experience(experience_text or text)
    if total_years > 40:
        warning_flags.append(f"Unrealistic Experience: Candidate claims {total_years} years of work experience.")
        fraud_score += 30

    # Specific technology release date checks
    tech_release_checks = [
        ("docker", 2013, 13),
        ("kubernetes", 2014, 12),
        ("react", 2013, 13),
        ("angular", 2010, 16),
        ("vue", 2014, 12),
        ("tensorflow", 2015, 11),
    ]
    for tech, release_year, max_years in tech_release_checks:
        if tech in text_lower:
            # Look for experience matching tech
            pattern = r'(?:experience(?:\s+with|\s+in)?\s+|\b' + tech + r'\b[^\d]*?)(\d+)\+?\s*(?:years|yrs)'
            matches = re.findall(pattern, text_lower)
            for match in matches:
                years = int(match)
                if years > max_years:
                    warning_flags.append(
                        f"Unrealistic Experience: Claims {years} years of {tech.title()} experience "
                        f"({tech.title()} was released in {release_year}, max {max_years} years)."
                    )
                    fraud_score += 25

    # 3. Suspicious Timeline Gaps (Work Gaps)
    exp_to_search = experience_text or text
    years = sorted(list(set([int(y) for y in re.findall(r'\b(19\d{2}|20\d{2})\b', exp_to_search)])))
    # Check gap between consecutive years in experience
    if len(years) >= 2:
        max_gap = 0
        for i in range(len(years) - 1):
            gap = years[i+1] - years[i]
            # Ignore gaps that look like graduation-to-work or ancient gaps
            if gap > 3 and years[i] > 2005:
                max_gap = max(max_gap, gap)
        if max_gap > 3:
            warning_flags.append(f"Suspicious Timeline Gap: A career gap of {max_gap} years was detected in history.")
            fraud_score += 15

    # 4. Hidden Text / Skill Padding Density Check
    # Look for long blocks of text with abnormally high skill density
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    for p in paragraphs[:15]: # check top paragraphs where resume keywords are stuffed
        if len(p) > 100:
            p_skills = extract_skills(p)
            if p_skills:
                # Calculate what fraction of characters are within skill words
                skill_chars = sum(len(s) for s in p_skills)
                density = skill_chars / len(p)
                if density > 0.65 and len(p_skills) > 8:
                    warning_flags.append("Hidden Text/Skill Padding: Extremely high keyword density (skills block with no prose) detected.")
                    fraud_score += 20
                    break

    # Cap fraud score at 100
    fraud_score = min(fraud_score, 100)
    
    return fraud_score, warning_flags

